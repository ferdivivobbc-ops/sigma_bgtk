import os
import io
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse, Http404
from django.conf import settings
from datetime import date

from accounts.decorators import role_required, mahasiswa_required
from mahasiswa.models import Mahasiswa
from magang.models import Magang
from absensi.models import Absensi
from kegiatan.models import KegiatanHarian
from laporan.models import LaporanAkhir
from accounts.models import Notification, AuditLog

# PDF Generation imports (ReportLab)
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import cm

# Word Generation imports (python-docx)
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _can_access_laporan(user, laporan):
    if user.is_admin():
        return True
    if user.is_mahasiswa():
        return laporan.mahasiswa.user_id == user.id
    if user.is_pembimbing():
        magang = getattr(laporan.mahasiswa, 'data_magang', None)
        return magang is not None and magang.pembimbing_id == user.id
    return False

@login_required
def laporan_detail(request, pk=None):
    user = request.user
    if user.is_mahasiswa():
        mhs = user.mahasiswa_profile
        laporan = LaporanAkhir.objects.filter(mahasiswa=mhs).first()
    else:
        if pk:
            laporan = get_object_or_404(LaporanAkhir.objects.select_related('mahasiswa__user', 'mahasiswa__kampus'), pk=pk)
        else:
            laporan = LaporanAkhir.objects.select_related('mahasiswa__user', 'mahasiswa__kampus').first()

        if laporan and not _can_access_laporan(user, laporan):
            messages.error(request, "Anda tidak memiliki akses ke laporan tersebut.")
            return redirect('dashboard')

    if not laporan and user.is_mahasiswa():
        return render(request, 'laporan/laporan_detail.html', {'laporan': None, 'no_laporan': True})

    return render(request, 'laporan/laporan_detail.html', {'laporan': laporan})

@login_required
@mahasiswa_required
def laporan_generate_draft(request):
    mhs = request.user.mahasiswa_profile
    magang = getattr(mhs, 'data_magang', None)

    if not magang:
        messages.error(request, "Anda belum memiliki data magang aktif untuk membuat laporan akhir.")
        return redirect('dashboard')

    total_hadir = Absensi.objects.filter(mahasiswa=mhs, status='HADIR').count()
    kegiatan_approved = KegiatanHarian.objects.filter(mahasiswa=mhs, status_verifikasi='DISETUJUI')
    total_kegiatan = kegiatan_approved.count()

    judul_default = f"LAPORAN AKHIR PRAKTEK KERJA LAPANGAN / MAGANG DI BGTK SUMATERA BARAT"
    kata_pengantar_default = (
        f"Puji dan syukur penulis panjatkan kehadirat Allah SWT atas rahmat dan karunia-Nya "
        f"sehingga penulis dapat menyelesaikan kegiatan magang serta menyusun Laporan Akhir Magang ini dengan baik.\n\n"
        f"Laporan ini disusun sebagai bentuk pertanggungjawaban atas pelaksanaan magang yang dilaksanakan pada unit "
        f"{magang.bagian} Balai Guru dan Tenaga Kependidikan (BGTK) Sumatera Barat dari tanggal {magang.tanggal_mulai} "
        f"sampai {magang.tanggal_selesai}.\n\n"
        f"Penulis mengucapkan terima kasih kepada Pembimbing BGTK {magang.pembimbing.get_full_name() if magang.pembimbing else 'BGTK Sumbar'}, "
        f"serta seluruh pihak kampus {mhs.kampus.nama_kampus if mhs.kampus else 'kampus asal'} yang telah memberikan bimbingan dan arahan."
    )

    bab1_latar_belakang = (
        "Kegiatan magang merupakan salah satu sarana penting bagi mahasiswa untuk menerapkan ilmu pengetahuan akademis "
        "ke dalam dunia kerja nyata. BGTK Sumatera Barat sebagai instansi yang membina guru dan tenaga kependidikan "
        "memberikan kesempatan emas bagi mahasiswa untuk memahami tata kelola administrasi dan operasional pemerintahan."
    )
    bab1_tujuan = (
        "1. Memahami proses kerja dan administrasi di BGTK Sumatera Barat.\n"
        "2. Meningkatkan keterampilan praktis dan profesionalisme dalam bidang studi yang ditekuni.\n"
        "3. Memenuhi salah satu syarat kelulusan program magang perguruan tinggi."
    )
    bab1_manfaat = (
        "1. Bagi Mahasiswa: Memperoleh pengalaman kerja nyata dan wawasan operasional di instansi pemerintahan.\n"
        "2. Bagi Perguruan Tinggi: Terjalinnya hubungan kerjasama yang erat dengan BGTK Sumatera Barat.\n"
        "3. Bagi BGTK Sumatera Barat: Terbantunya pelaksanaan administrasi dan kegiatan operasional harian."
    )

    bab2_gambaran_umum = (
        f"Balai Guru dan Tenaga Kependidikan (BGTK) Sumatera Barat merupakan unit pelaksana teknis yang bertugas "
        f"melaksanakan pengembangan dan pemberdayaan guru, pendidik lainnya, dan tenaga kependidikan.\n\n"
        f"Selama masa magang, mahasiswa ditempatkan pada bagian: {magang.bagian} di bawah bimbingan "
        f"{magang.pembimbing.get_full_name() if magang.pembimbing else 'Pembimbing BGTK'}."
    )

    activities_text = "\n".join([f"- {k.tanggal.strftime('%d-%m-%Y')}: {k.judul_kegiatan} - {k.hasil}" for k in kegiatan_approved[:8]])
    bab3_pelaksanaan = (
        f"Pelaksanaan magang berlangsung mulai tanggal {magang.tanggal_mulai} hingga {magang.tanggal_selesai}.\n"
        f"Rekapitulasi Kehadiran: Total hadir sebanyak {total_hadir} hari kerja.\n"
        f"Ringkasan Kegiatan Utama yang dilaksanakan:\n{activities_text or 'Belum ada kegiatan yang disetujui.'}"
    )

    bab4_hasil = (
        f"Selama masa magang, mahasiswa berhasil menyelesaikan {total_kegiatan} kegiatan yang terverifikasi.\n\n"
        "Kompetensi yang diperoleh meliputi pengelolaan data administrasi, koordinasi kerja tim, penyelesaian masalah harian, "
        "dan penerapan etika kerja profesional."
    )

    bab5_kesimpulan = (
        "1. Kegiatan magang di BGTK Sumatera Barat berjalan dengan lancar dan memberikan wawasan praktis yang sangat berharga.\n"
        "2. Seluruh target tugas harian dan administrasi dapat diselesaikan dengan baik sesuai arahan pembimbing."
    )
    bab5_saran = (
        "1. Untuk BGTK Sumatera Barat: Diharapkan terus mempertahankan sistem penerimaan magang yang terstruktur.\n"
        "2. Untuk Mahasiswa Selanjutnya: Agar selalu aktif berkomunikasi dan menjaga kedisiplinan selama masa magang."
    )

    laporan, created = LaporanAkhir.objects.get_or_create(
        mahasiswa=mhs,
        defaults={
            'judul': judul_default,
            'kata_pengantar': kata_pengantar_default,
            'latar_belakang': bab1_latar_belakang,
            'tujuan': bab1_tujuan,
            'manfaat': bab1_manfaat,
            'gambaran_umum': bab2_gambaran_umum,
            'pelaksanaan_magang': bab3_pelaksanaan,
            'hasil_pembahasan': bab4_hasil,
            'kesimpulan': bab5_kesimpulan,
            'saran': bab5_saran,
            'status': 'DRAFT',
        }
    )

    if created:
        messages.success(request, "Draft Laporan Akhir berhasil digenerate dari data magang, absensi, dan kegiatan harian Anda.")
    else:
        messages.info(request, "Draft laporan sudah tersedia. Isi laporan yang telah Anda edit tetap dipertahankan.")
    return redirect('laporan_detail_pk', pk=laporan.pk)

@login_required
@mahasiswa_required
def laporan_edit(request, pk):
    laporan = get_object_or_404(LaporanAkhir, pk=pk, mahasiswa=request.user.mahasiswa_profile)

    if request.method == 'POST':
        laporan.judul = request.POST.get('judul', '').strip()
        laporan.kata_pengantar = request.POST.get('kata_pengantar', '').strip()
        laporan.latar_belakang = request.POST.get('latar_belakang', '').strip()
        laporan.tujuan = request.POST.get('tujuan', '').strip()
        laporan.manfaat = request.POST.get('manfaat', '').strip()
        laporan.gambaran_umum = request.POST.get('gambaran_umum', '').strip()
        laporan.pelaksanaan_magang = request.POST.get('pelaksanaan_magang', '').strip()
        laporan.hasil_pembahasan = request.POST.get('hasil_pembahasan', '').strip()
        laporan.kesimpulan = request.POST.get('kesimpulan', '').strip()
        laporan.saran = request.POST.get('saran', '').strip()

        uploaded_pdf = request.FILES.get('file_pdf')
        uploaded_word = request.FILES.get('file_word')
        upload_errors = []

        for uploaded_file, allowed_extensions, label in (
            (uploaded_pdf, {'.pdf'}, 'File PDF'),
            (uploaded_word, {'.doc', '.docx'}, 'File Word'),
        ):
            if uploaded_file:
                extension = os.path.splitext(uploaded_file.name)[1].lower()
                if extension not in allowed_extensions:
                    upload_errors.append(f'{label} harus berformat {", ".join(sorted(allowed_extensions))}.')
                elif uploaded_file.size > 10 * 1024 * 1024:
                    upload_errors.append(f'{label} tidak boleh lebih dari 10 MB.')

        if upload_errors:
            for error in upload_errors:
                messages.error(request, error)
            return render(request, 'laporan/laporan_form.html', {'laporan': laporan})

        if uploaded_pdf:
            laporan.file_pdf = uploaded_pdf
        if uploaded_word:
            laporan.file_word = uploaded_word

        if request.POST.get('submit_approval'):
            laporan.status = 'PERIKSA'
            messages.success(request, "Laporan akhir berhasil diperbarui dan diajukan ke pembimbing untuk diperiksa.")
        else:
            messages.success(request, "Draft laporan akhir berhasil diperbarui.")

        laporan.save()
        AuditLog.objects.create(actor=request.user, action='SUBMIT' if request.POST.get('submit_approval') else 'UPDATE', model_name='LaporanAkhir', object_id=laporan.pk, description=laporan.status)
        if request.POST.get('submit_approval'):
            magang = getattr(laporan.mahasiswa, 'data_magang', None)
            if magang and magang.pembimbing_id:
                Notification.objects.create(
                    recipient=magang.pembimbing,
                    title='Laporan akhir menunggu pemeriksaan',
                    message=f'{laporan.mahasiswa.user.get_full_name()} mengajukan laporan akhir untuk diperiksa.',
                    link=f'/laporan/{laporan.pk}/',
                )
        return redirect('laporan_detail_pk', pk=laporan.pk)

    return render(request, 'laporan/laporan_form.html', {'laporan': laporan})

@login_required
@role_required('ADMIN', 'PEMBIMBING')
def laporan_status_update(request, pk):
    laporan = get_object_or_404(LaporanAkhir, pk=pk)

    if not _can_access_laporan(request.user, laporan):
        messages.error(request, "Anda tidak memiliki akses ke laporan tersebut.")
        return redirect('dashboard')

    if request.method == 'POST':
        status = request.POST.get('status', 'DISETUJUI')
        catatan = request.POST.get('catatan_pembimbing', '').strip()

        laporan.status = status
        laporan.catatan_pembimbing = catatan
        laporan.save()
        AuditLog.objects.create(actor=request.user, action='STATUS', model_name='LaporanAkhir', object_id=laporan.pk, description=status)
        Notification.objects.create(
            recipient=laporan.mahasiswa.user,
            title='Status laporan akhir diperbarui',
            message=f'Laporan akhir Anda berstatus {laporan.get_status_display()}.',
            link=f'/laporan/{laporan.pk}/',
        )

        messages.success(request, f"Status Laporan Akhir {laporan.mahasiswa.user.get_full_name()} diubah menjadi '{laporan.get_status_display()}'.")
        return redirect('laporan_detail_pk', pk=laporan.pk)

    return redirect('laporan_detail_pk', pk=laporan.pk)

@login_required
def export_pdf(request, pk):
    laporan = get_object_or_404(LaporanAkhir.objects.select_related('mahasiswa__user', 'mahasiswa__kampus'), pk=pk)
    if not _can_access_laporan(request.user, laporan):
        messages.error(request, "Anda tidak memiliki akses ke laporan tersebut.")
        return redirect('dashboard')
    mhs = laporan.mahasiswa
    magang = getattr(mhs, 'data_magang', None)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CoverTitle',
        parent=styles['Heading1'],
        alignment=1, # Center
        fontSize=16,
        leading=22,
        textColor=colors.HexColor('#0F172A'),
        spaceAfter=15
    )

    header_style = ParagraphStyle(
        'SectionHeader',
        parent=styles['Heading2'],
        fontSize=13,
        leading=16,
        textColor=colors.HexColor('#1E3A8A'),
        spaceBefore=15,
        spaceAfter=8
    )

    body_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontSize=10,
        leading=14,
        spaceAfter=6,
        alignment=4 # Justify
    )

    story = []

    # COVER PAGE
    story.append(Paragraph("SIGMA - BGTK SUMATERA BARAT", ParagraphStyle('Sub', alignment=1, fontSize=12, textColor=colors.HexColor('#475569'))))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(f"<b>{laporan.judul.upper()}</b>", title_style))
    story.append(Spacer(1, 1.5*cm))

    cover_info = [
        [Paragraph("<b>Nama Mahasiswa</b>", body_style), Paragraph(f": {mhs.user.get_full_name()}", body_style)],
        [Paragraph("<b>NIM</b>", body_style), Paragraph(f": {mhs.nim}", body_style)],
        [Paragraph("<b>Program Studi</b>", body_style), Paragraph(f": {mhs.prodi}", body_style)],
        [Paragraph("<b>Perguruan Tinggi</b>", body_style), Paragraph(f": {mhs.kampus.nama_kampus if mhs.kampus else '-'}", body_style)],
        [Paragraph("<b>Bagian / Unit BGTK</b>", body_style), Paragraph(f": {magang.bagian if magang else '-'}", body_style)],
        [Paragraph("<b>Pembimbing BGTK</b>", body_style), Paragraph(f": {magang.pembimbing.get_full_name() if magang and magang.pembimbing else '-'}", body_style)],
        [Paragraph("<b>Periode Magang</b>", body_style), Paragraph(f": {magang.tanggal_mulai if magang else '-'} s/d {magang.tanggal_selesai if magang else '-'}", body_style)],
    ]

    t_info = Table(cover_info, colWidths=[5*cm, 11*cm])
    t_info.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(t_info)
    story.append(Spacer(1, 2*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#CBD5E1'), spaceAfter=15))

    # KATA PENGANTAR
    story.append(Paragraph("<b>KATA PENGANTAR</b>", header_style))
    story.append(Paragraph(laporan.kata_pengantar or "-", body_style))
    story.append(Spacer(1, 10))

    # BAB I
    story.append(Paragraph("<b>BAB I – PENDAHULUAN</b>", header_style))
    story.append(Paragraph("<b>1.1 Latar Belakang</b>", ParagraphStyle('SubSub', parent=body_style, fontSize=11)))
    story.append(Paragraph(laporan.latar_belakang or "-", body_style))
    story.append(Paragraph("<b>1.2 Tujuan Magang</b>", ParagraphStyle('SubSub', parent=body_style, fontSize=11)))
    story.append(Paragraph(laporan.tujuan or "-", body_style))
    story.append(Paragraph("<b>1.3 Manfaat Magang</b>", ParagraphStyle('SubSub', parent=body_style, fontSize=11)))
    story.append(Paragraph(laporan.manfaat or "-", body_style))
    story.append(Spacer(1, 10))

    # BAB II
    story.append(Paragraph("<b>BAB II – GAMBARAN UMUM BGTK SUMATERA BARAT</b>", header_style))
    story.append(Paragraph(laporan.gambaran_umum or "-", body_style))
    story.append(Spacer(1, 10))

    # BAB III
    story.append(Paragraph("<b>BAB III – PELAKSANAAN MAGANG</b>", header_style))
    story.append(Paragraph(laporan.pelaksanaan_magang or "-", body_style))
    story.append(Spacer(1, 10))

    # BAB IV
    story.append(Paragraph("<b>BAB IV – HASIL DAN PEMBAHASAN</b>", header_style))
    story.append(Paragraph(laporan.hasil_pembahasan or "-", body_style))
    story.append(Spacer(1, 10))

    # BAB V
    story.append(Paragraph("<b>BAB V – PENUTUP</b>", header_style))
    story.append(Paragraph("<b>5.1 Kesimpulan</b>", ParagraphStyle('SubSub', parent=body_style, fontSize=11)))
    story.append(Paragraph(laporan.kesimpulan or "-", body_style))
    story.append(Paragraph("<b>5.2 Saran</b>", ParagraphStyle('SubSub', parent=body_style, fontSize=11)))
    story.append(Paragraph(laporan.saran or "-", body_style))
    story.append(Spacer(1, 15))

    # LAMPIRAN KEGIATAN HARIAN
    story.append(Paragraph("<b>LAMPIRAN – REKAP KEGIATAN HARIAN TERVERIFIKASI</b>", header_style))
    kegiatan_list = KegiatanHarian.objects.filter(mahasiswa=mhs, status_verifikasi='DISETUJUI').order_by('tanggal')[:10]

    if kegiatan_list:
        table_data = [["No", "Tanggal", "Judul Kegiatan", "Hasil Kegiatan"]]
        for idx, k in enumerate(kegiatan_list, 1):
            table_data.append([
                str(idx),
                k.tanggal.strftime('%d/%m/%Y'),
                Paragraph(k.judul_kegiatan, body_style),
                Paragraph(k.hasil, body_style)
            ])

        t_kegiatan = Table(table_data, colWidths=[1*cm, 3*cm, 6*cm, 6*cm])
        t_kegiatan.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1E3A8A')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
            ('ALIGN', (0,0), (1,-1), 'CENTER'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#94A3B8')),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
        ]))
        story.append(t_kegiatan)
    else:
        story.append(Paragraph("Belum ada kegiatan harian terverifikasi.", body_style))

    doc.build(story)
    pdf_value = buffer.getvalue()
    buffer.close()

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="Laporan_Akhir_{mhs.nim}.pdf"'
    response.write(pdf_value)
    return response

@login_required
def export_word(request, pk):
    laporan = get_object_or_404(LaporanAkhir.objects.select_related('mahasiswa__user', 'mahasiswa__kampus'), pk=pk)
    if not _can_access_laporan(request.user, laporan):
        messages.error(request, "Anda tidak memiliki akses ke laporan tersebut.")
        return redirect('dashboard')
    mhs = laporan.mahasiswa
    magang = getattr(mhs, 'data_magang', None)

    doc = docx.Document()

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("SIGMA - BGTK SUMATERA BARAT\n\n" + laporan.judul.upper())
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_paragraph() # spacing

    # Student Info Table
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    infos = [
        ("Nama Mahasiswa", mhs.user.get_full_name()),
        ("NIM", mhs.nim),
        ("Program Studi", mhs.prodi),
        ("Kampus", mhs.kampus.nama_kampus if mhs.kampus else "-"),
        ("Bagian BGTK", magang.bagian if magang else "-"),
        ("Pembimbing BGTK", magang.pembimbing.get_full_name() if magang and magang.pembimbing else "-"),
    ]
    for label, val in infos:
        row_cells = info_table.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = f": {val}"

    doc.add_page_break()

    # Sections
    sections = [
        ("KATA PENGANTAR", laporan.kata_pengantar),
        ("BAB I – PENDAHULUAN", f"1.1 Latar Belakang\n{laporan.latar_belakang}\n\n1.2 Tujuan\n{laporan.tujuan}\n\n1.3 Manfaat\n{laporan.manfaat}"),
        ("BAB II – GAMBARAN UMUM BGTK SUMBAR", laporan.gambaran_umum),
        ("BAB III – PELAKSANAAN MAGANG", laporan.pelaksanaan_magang),
        ("BAB IV – HASIL DAN PEMBAHASAN", laporan.hasil_pembahasan),
        ("BAB V – PENUTUP", f"5.1 Kesimpulan\n{laporan.kesimpulan}\n\n5.2 Saran\n{laporan.saran}"),
    ]

    for title, content in sections:
        h = doc.add_heading(title, level=1)
        doc.add_paragraph(content or "-")
        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    word_value = buffer.getvalue()
    buffer.close()

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Laporan_Akhir_{mhs.nim}.docx"'
    response.write(word_value)
    return response
